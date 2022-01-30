from fileinput import close
import os
import sys
import subprocess
import hilight
import docx
from docx.shared import RGBColor
from html.parser import HTMLParser
try:
	from pygments import lexers
	from pygments.formatters import HtmlFormatter
	from pygments import highlight
except:
	print('Error: you need to install pygments use \'pip install pygments\'')
	exit(0)
try:
  import docx
  from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
  print('Error: you need to install python-docx use \'pip install python-docx\'')
  exit(0)
try:
  import reportlab
  from reportlab.lib.units import mm
  from reportlab.pdfgen import canvas
  from reportlab.lib.pagesizes import A4
  from reportlab.pdfbase.pdfmetrics import stringWidth
except ImportError:
  print('Error: you need to install reportlab use \'pip install reportlab\'')
  exit(0)
try:
  from PyPDF2 import PdfFileWriter, PdfFileReader
except ImportError:
  print('Error: you need to install reportlab use \'pip install reportlab\'')
  exit(0)

# try:
#     import pdftotext
# except ImportError:
#     print('''Error: you need to install pdftotext use 'pip install pdftotext' if it didn't 
#   work go to https://github.com/jalan/pdftotext to see how to install it
#   also Microsoft Visual C++ 14.0 or greater is required''')
#     exit(0)

wDir = os.getcwd()
if len(sys.argv) > 1:
    wDir = ' '.join(sys.argv[1:])
doc = docx.Document()
section = doc.sections[0]
# set size to A4
section.page_height = docx.shared.Mm(297)
section.page_width = docx.shared.Mm(210)
# set margins
section.top_margin = docx.shared.Mm(20)
section.bottom_margin = docx.shared.Mm(20)
section.left_margin = docx.shared.Mm(20)
section.right_margin = docx.shared.Mm(15)
pdf = None


def generate_docx():
    # generates the initial doc file
    for d in os.listdir(wDir):
        if os.path.isfile(d):
            continue
        doc.add_heading(d, 0)
        for f in os.listdir(d):
            file = os.path.join(os.path.abspath(wDir), d, f)
            add_file(file)
    doc.save('x.docx')


def convert_to_pdf(x):
    # convert docx to pdf
    # x is the file name

    # if you are a linux user try this :D
    # try:
    #     c = 'lowriter --convert-to pdf '+x+'.docx'
    #     cp = subprocess.run([c], shell=True, universal_newlines=True,
    #                         stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    #     # if len(cp.stderr):
    #     #   raise Exception('libre office')
    # except:
    #     print('Error: cannot find instance from libre office')
    #     exit(0)


    from docx2pdf import convert
    convert(x+".docx")

    # import aspose.words as aw
    # doc = aw.Document(x+".docx")
    # doc.save("x.pdf")

    os.remove(x+'.docx')


def createPagePdf():
    # generate new pdf contains pages numbers which will be merged with origenal pdf later
    c = canvas.Canvas('num.pdf', pagesize=A4)
    num = pdf.getNumPages()
    for i in range(1, num+1):
        c.drawString((210//2)*mm, (8)*mm, str(i))
        c.showPage()
    c.save()


def add_pages_num():
    # merge pdf and num pdf
    # to add page numbers
    npdf = open('num.pdf', 'rb')
    numpdf = PdfFileReader(npdf)
    pdfWriter = PdfFileWriter()
    for p in range(pdf.getNumPages()):
        a = pdf.getPage(p)
        a.mergePage(numpdf.getPage(p))
        pdfWriter.addPage(a)
    res = open('z.pdf', 'wb')
    pdfWriter.write(res)
    res.close()
    npdf.close()
    os.remove('num.pdf')


def create_content():
    # create table of content for your document
    # pp = pdftotext.PDF(open('l.pdf', 'rb'))
    pdffileobj=open('z.pdf','rb')
    pdfreader=PdfFileReader(pdffileobj)
    x=pdfreader.numPages
    pp = []
    for i in range(x):
        pp.append(pdfreader.getPage(i).extractText().replace('\n', ''))
    c = canvas.Canvas('pre.pdf', pagesize=A4)
    pageNum = 0
    width, height = A4
    currH = 760
    c.setFont('Helvetica', 20)
    sw = stringWidth('Content', 'Helvetica', 20)
    c.drawString(width/2-sw/2, 800, 'Content')
    for d in os.listdir(wDir):
        if os.path.isfile(d):
            continue
        if d == 'l.pdf':
            continue
        first = True
        for f in os.listdir(d):
            f = '.'.join(f.split('.')[:-1])
            f = f.strip()
            while(f not in pp[pageNum]):
                pageNum = pageNum + 1
            if first:
                currH -= 20
                if currH < 100:
                    c.showPage()
                    currH = 760
                c.setFont('Helvetica-Bold', 16)
                c.drawString(40, currH, d)
                currH -= 50
                first = False
            if currH < 50:
                c.showPage()
                currH = 760
            c.setFont('Helvetica', 12)
            c.drawString(40, currH, f)
            dotcz = stringWidth('.', 'Helvetica', 12)
            curcz = 45 + stringWidth(f, 'Helvetica', 12)
            c.drawString(curcz, currH, int((width-65-curcz)//dotcz)*'.')
            c.drawString(width-60, currH, str(pageNum+1))
            currH -= 30
    c.showPage()
    c.save()
    pdffileobj.close()


def merge_content():
    # merges pre.pdf(content pdf) with l.pdf(main pdf)
    writer = PdfFileWriter()
    prepdf = open('pre.pdf', 'rb')
    contnt = PdfFileReader(prepdf)
    for p in range(contnt.getNumPages()):
        writer.addPage(contnt.getPage(p))
    for p in range(pdf.getNumPages()):
        writer.addPage(pdf.getPage(p))
    res = open('lib.pdf', 'wb')
    writer.write(res)
    prepdf.close()
    os.remove('pre.pdf')


def add_file(f):
    # Just adds algorithm to the document
    file = open(f, 'r')
    algo = os.path.split(f)[1]  # get file name without abs path
    algo = '.'.join(algo.split('.')[:-1])  # get name without extension
    algo = algo.strip()
    doc.add_paragraph(algo, style='List Bullet')
    doc.paragraphs[-1].runs[0].font.size = docx.shared.Pt(16)
    doc.paragraphs[-1].runs[0].bold = True
    code = ''
    for line in file:
        code += line
    lang = f.split('.')[-1]
    if(lang == 'txt'):
        lang = 'cpp'
    hilight.hilight(code, lang)
    hilight.generateDoc()
    tmp = docx.Document('st.docx')
    doc.add_paragraph()
    for run in tmp.paragraphs[0].runs:
        doc.paragraphs[-1].add_run(run.text)
        doc.paragraphs[-1].runs[-1].bold = run.bold
        doc.paragraphs[-1].runs[-1].font.color.rgb = run.font.color.rgb
    os.remove('st.docx')


def main():
    global pdf
    print('creating initial doc..', end=' ')
    generate_docx()
    print('done\ncreating pdf.. done')
    convert_to_pdf('x')
    xpdf = open('x.pdf', 'rb')
    pdf = PdfFileReader(xpdf)
    createPagePdf()
    print('adding pages numbers..', end=' ')
    add_pages_num()
    zpdf = open('z.pdf', 'rb')
    pdf = PdfFileReader(zpdf)
    print('done\nadding content pages..', end=' ')
    create_content()
    print('done\nmerging content with origenal file..', end=' ')
    merge_content()
    print('done\nyour file is ready now :D')
    zpdf.close()
    xpdf.close()
    os.remove('z.pdf')
    os.remove('x.pdf')


if __name__ == '__main__':
    main()
