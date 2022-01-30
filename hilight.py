import os
import docx
from docx.shared import RGBColor
from html.parser import HTMLParser
try:
	from pygments import lexers
	from pygments.formatters import HtmlFormatter
	from pygments import highlight
except:
	print('Error: you need to install pygments use \'pip install pygments\'')
	exit(0)


html_parser = HTMLParser()

def hilight(text, lang):
	# generate html and css files for highlighted code
	lex = lexers.get_lexer_by_name(lang)
	formatter = HtmlFormatter(style='colorful') # linenos=True for line numbers
	with open('out.html', 'w') as f:
		highlight(text, lex, formatter, outfile=f)
	with open('out.html', 'r') as f:
		lines = f.readlines()
	with open('out.html', 'w') as f:
		f.write('<link rel="stylesheet" type="text/css" href="style.css">\n')
		for line in lines:
			f.write(line)
	with open('style.css', 'w') as f:
    		f.write(formatter.get_style_defs())


doc = docx.Document()
doc.add_paragraph()

styles = {}

def addToFDoc(txt, _class):
	# add txt with the styles extracted from css file using _class
	doc.paragraphs[0].add_run(txt)
	p = ''
	if _class in styles:
		p = styles[_class]
	idx = p.find('color:')
	if idx != -1:
		idx += 8
		col = p[idx:idx+6]
		doc.paragraphs[0].runs[-1].font.color.rgb = RGBColor(int(col[:2], 16), int(col[2:4], 16), int(col[4:], 16))
	idx = p.find('font-weight: bold')
	if idx != -1:
		doc.paragraphs[0].runs[-1].bold = True


def generateDoc():
	# get test and class to be added in the doc
	
	# add styles to dectunary to make it easier to get style for objects
	global styles, doc
	styles = {}
	doc = docx.Document()
	doc.add_paragraph()
	with open('style.css', 'r') as c:
		for line in c.readlines():
			styles[line[1:line.find(' ')]] = line

	with open('out.html', 'r') as h:
		ht = h.read()
		lst = 0
		while True:
			idx = ht.find('span class="', lst)
			lst = idx+12
			if idx == -1:
				break
			idx = ht.find('"', lst)
			_class = ht[lst:idx]
			lst = idx+2
			idx = ht.find('</span>', lst)
			txt = ht[lst:idx]
			lst = idx+7
			idx = ht.find('<', lst)
			txt = html_parser.unescape(txt)
			addToFDoc(txt, _class)
			doc.paragraphs[0].add_run(ht[lst:idx])
			# print(txt, _class)
	doc.save('st.docx')
	os.remove('out.html')
	os.remove('style.css')


if __name__ == '__main__':
	hilight('''#include<bits/stdc++.h>
	using namespace std;

	int n;
	int main(){
		ios::sync_with_stdio(0); cin.tie(0); cout.tie(0);
	#ifndef ONLINE_JUDGE
	//  freopen("in.txt", "r", stdin);
	#endif // ONLINE_JUDGE
		cin >> n;
		if(n&1) cout << 0 << '\\n';
		else cout << (1LL<<(n/2)) << '\\n';
		return 0;
	}''', 'cpp')
	generateDoc()
