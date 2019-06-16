import os
import sys
import subprocess
try:
  import docx
  from docx.enum.text import WD_ALIGN_PARAGRAPH 
except ImportError:
  print('Error: you need to install python-docx use \'pip install python-docx\'')
  exit(0)
try:
  import reportlab
  from reportlab.lib.units import mm
  from reportlab.pdfgen import canvas
  from reportlab.lib.pagesizes import A4
  from reportlab.pdfbase.pdfmetrics import stringWidth
except ImportError:
  print('Error: you need to install reportlab use \'pip install reportlab\'')
  exit(0)
try:
  from PyPDF2 import PdfFileWriter, PdfFileReader
except ImportError:
  print('Error: you need to install reportlab use \'pip install reportlab\'')
  exit(0)
try:
  import pdftotext
except ImportError:
  print('''Error: you need to install pdftotext use 'pip install pdftotext' if it didn't 
  work go to https://github.com/jalan/pdftotext to see how to install it''')
  exit(0)

wDir = os.getcwd()
if len(sys.argv) > 1:
  wDir = ' '.join(sys.argv[1:])
doc = docx.Document()
section = doc.sections[0]
# set size to A4
section.page_height = docx.shared.Mm(297)
section.page_width = docx.shared.Mm(210)
# set margins
section.top_margin = docx.shared.Mm(20)
section.bottom_margin = docx.shared.Mm(20)
section.left_margin = docx.shared.Mm(20)
section.right_margin = docx.shared.Mm(15)
pdf = None
numpdf = None


def generate_docx():
  # generates the initial doc file
  for d in os.listdir(wDir):
    doc.add_heading(d, 0)
    for f in os.listdir(d):
      file = os.path.join(os.path.abspath(wDir), d, f)
      add_file(file)
  doc.save('x.docx')



def convert_to_pdf(x):
  try:
    c = 'lowriter --convert-to pdf '+x+'.docx'
    cp = subprocess.run([c], shell=True, universal_newlines=True,
     stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    # if len(cp.stderr):
    #   raise Exception('libre office')
  except:
    print('Error: cannot find instance from libre office')
    exit(0)
  os.remove(x+'.docx')


def createPagePdf():
  # generate new pdf contains pages numbers which we will merge with origenal pdf later
  c = canvas.Canvas('num.pdf', pagesize = A4)
  num = pdf.getNumPages()
  for i in range(1,num+1): 
    c.drawString((210//2)*mm, (8)*mm, str(i))
    c.showPage()
  c.save()
  global numpdf
  numpdf = PdfFileReader(open('num.pdf', 'rb'))


def add_pages_num():
  # merge pdf and num pdf 
  pdfWriter = PdfFileWriter()
  for p in range(pdf.getNumPages()):
    a = pdf.getPage(p)
    a.mergePage(numpdf.getPage(p))
    pdfWriter.addPage(a)
  res = open('l.pdf', 'wb')
  pdfWriter.write(res)
  os.remove('x.pdf')
  os.remove('num.pdf')


def create_content():
  # create table of content for your document
  pp = pdftotext.PDF(open('l.pdf', 'rb'))
  c = canvas.Canvas('pre.pdf', pagesize = A4)
  pageNum = 0
  width, height = A4
  currH = 760
  c.setFont('Helvetica', 20)
  sw = stringWidth('Content', 'Helvetica', 20)
  c.drawString(width/2-sw/2, 800, 'Content')
  for d in os.listdir(wDir):
    if d == 'l.pdf':
      continue
    first = True
    for f in os.listdir(d):
      f = '.'.join(f.split('.')[:-1])
      while(f not in pp[pageNum]):
        pageNum = pageNum + 1
      if first:
        currH -= 20
        if currH < 100:
          c.showPage()
          currH = 760
        c.setFont('Helvetica', 16)
        c.drawString(40, currH, d)
        currH -= 50
        first = False
      if currH < 50:
        c.showPage()
        currH = 760
      c.setFont('Helvetica', 12)
      c.drawString(40, currH, f)
      dotcz = stringWidth('.', 'Helvetica', 12)
      curcz = 45 + stringWidth(f, 'Helvetica', 12) 
      c.drawString(curcz, currH, int((width-65-curcz)//dotcz)*'.')
      c.drawString(width-60, currH, str(pageNum+1))
      currH -= 30
  c.showPage()
  c.save()



def merge_content():
  # merges pre.pdf(content pdf) with l.pdf(main pdf)
  writer = PdfFileWriter()
  contnt = PdfFileReader(open('pre.pdf', 'rb'))
  for p in range(contnt.getNumPages()):
    writer.addPage(contnt.getPage(p))
  for p in range(pdf.getNumPages()):
    writer.addPage(pdf.getPage(p))
  res = open('lib.pdf', 'wb')
  writer.write(res)
  os.remove('l.pdf')
  os.remove('pre.pdf')

def add_file(f):
  file = open(f, 'r')
  algo = os.path.split(f)[1]  # get file name without abs path
  algo = '.'.join(algo.split('.')[:-1])  # get name without extension
  doc.add_paragraph(algo, style='List Bullet')
  doc.paragraphs[-1].runs[0].font.size = docx.shared.Pt(16)
  doc.paragraphs[-1].runs[0].bold = True
  code = ''
  for line in file:
    code += line
  doc.add_paragraph(code)


def main():
  global pdf
  print('creating initial doc..', end = ' ')
  generate_docx()
  print('done\ncreating pdf.. done')
  convert_to_pdf('x')
  pdf = PdfFileReader(open('x.pdf', 'rb'))
  createPagePdf()
  print('adding pages numbers..', end = ' ')
  add_pages_num()
  pdf = PdfFileReader(open('l.pdf', 'rb'))
  print('done\nadding content pages..', end = ' ')
  create_content()
  print('done\nmerging content with origenal file..', end = ' ')
  merge_content()
  print('done\nyour file is ready now :D')


if __name__ == '__main__':
  main()
